from __future__ import annotations

import asyncio
import io
import json
import zipfile
from typing import Any

import pytest
from docx import Document
from fastapi.testclient import TestClient

import app.main as main


@pytest.fixture
def client() -> TestClient:
    return TestClient(main.app)


def _result(text: str) -> dict[str, Any]:
    return {
        "corrected_text": text,
        "changes": [],
        "summary": "No issues.",
        "stats": {"words": 2, "report": {"findings": [], "scores": {"overall": 100}}},
    }


def test_health_root_and_static(client: TestClient) -> None:
    health = client.get("/health")
    assert health.status_code == 200
    assert health.json()["version"] == "3.0.0"
    assert health.json()["commit"] == "local"
    assert client.get("/").status_code == 200
    assert client.head("/").status_code == 200
    assert client.get("/static/index.html").status_code == 200


@pytest.mark.parametrize("text", ["", " \n "])
def test_proofread_rejects_empty(client: TestClient, text: str) -> None:
    assert client.post("/proofread", json={"text": text}).status_code == 400


def test_proofread_rejects_oversize(client: TestClient) -> None:
    response = client.post("/proofread", json={"text": "x" * 100_001})
    assert response.status_code == 413


def test_plain_proofread_and_history(
    client: TestClient, monkeypatch: pytest.MonkeyPatch
) -> None:
    received: dict[str, Any] = {}

    async def fake(
        text: str,
        progress: Any = None,
        status: Any = None,
        document_type: str = "auto",
        metadata: Any = None,
    ) -> dict[str, Any]:
        received.update({"type": document_type, "metadata": metadata})
        return _result(text)

    monkeypatch.setattr(main, "run_pipeline", fake)
    response = client.post(
        "/proofread",
        json={
            "text": "Hello world.",
            "document_type": "docx",
            "metadata": {"headings": []},
        },
    )
    assert response.status_code == 200
    job_id = response.json()["id"]
    assert client.get(f"/jobs/{job_id}").json()["original_text"] == "Hello world."
    assert client.get("/jobs?limit=1").json()[0]["id"] == job_id
    assert client.get("/jobs/not-found").status_code == 404
    assert received == {"type": "docx", "metadata": {"headings": []}}


def test_streaming_progress_result_and_headers(
    client: TestClient, monkeypatch: pytest.MonkeyPatch
) -> None:
    async def fake(
        text: str,
        progress: Any = None,
        status: Any = None,
        document_type: str = "auto",
        metadata: Any = None,
    ) -> dict[str, Any]:
        await progress(0, 1)
        await status("corrections", "Corrections", "running")
        await progress(1, 1)
        await status("corrections", "Corrections", "done")
        return _result(text)

    monkeypatch.setattr(main, "run_pipeline", fake)
    response = client.post("/proofread/stream", json={"text": "Hello world."})
    assert response.status_code == 200
    assert response.headers["content-type"].startswith("text/event-stream")
    assert response.headers["x-accel-buffering"] == "no"
    events = []
    for block in response.text.split("\n\n"):
        data = [
            line[5:].strip() for line in block.splitlines() if line.startswith("data:")
        ]
        if data:
            events.append(json.loads("\n".join(data)))
    assert [event["type"] for event in events] == [
        "progress",
        "agent",
        "progress",
        "agent",
        "result",
    ]


def test_streaming_sanitizes_internal_error(
    client: TestClient, monkeypatch: pytest.MonkeyPatch
) -> None:
    async def broken(*args: Any, **kwargs: Any) -> dict[str, Any]:
        raise ValueError("secret internal detail")

    monkeypatch.setattr(main, "run_pipeline", broken)
    response = client.post("/proofread/stream", json={"text": "Hello."})
    assert "secret internal detail" not in response.text
    assert '"type": "error"' in response.text


def test_plain_endpoint_converts_pipeline_timeout(
    client: TestClient, monkeypatch: pytest.MonkeyPatch
) -> None:
    async def slow(*args: Any, **kwargs: Any) -> dict[str, Any]:
        await asyncio.sleep(0.05)
        return _result("late")

    monkeypatch.setattr(main, "run_pipeline", slow)
    monkeypatch.setattr(main, "PIPELINE_TIMEOUT_SECONDS", 0.001)
    response = client.post("/proofread", json={"text": "Hello."})
    assert response.status_code == 503
    assert "timed out" in response.json()["detail"]


def test_text_markdown_and_unicode_extraction(client: TestClient) -> None:
    for filename, content in (
        ("sample.txt", "Café résumé 😀".encode()),
        ("sample.md", b"\xef\xbb\xbf# Heading\n\nText"),
    ):
        response = client.post("/extract", files={"file": (filename, content)})
        assert response.status_code == 200
        assert "\ufeff" not in response.json()["text"]
        assert response.json()["document_type"] in {"txt", "markdown"}


def test_docx_extraction_preserves_paragraph_table_order(client: TestClient) -> None:
    document = Document()
    document.add_heading("First 😀", level=1)
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    document.add_paragraph("Last")
    buffer = io.BytesIO()
    document.save(buffer)
    response = client.post(
        "/extract",
        files={"file": ("sample.docx", buffer.getvalue())},
    )
    assert response.status_code == 200
    assert response.json()["text"] == "First 😀\n\n| A | B |\n\nLast"
    assert response.json()["document_type"] == "docx"
    assert response.json()["metadata"]["headings"][0]["level"] == 1
    assert response.json()["metadata"]["tables"][0]["column_counts"] == [2]


def test_html_extraction_is_structured_and_ignores_scripts(client: TestClient) -> None:
    html = b"""
    <html><body><h1>Policy</h1><p>Useful text.</p>
    <script>ignore me</script><table><tr><td>A</td><td>B</td></tr></table>
    </body></html>
    """
    response = client.post("/extract", files={"file": ("policy.html", html)})
    assert response.status_code == 200
    payload = response.json()
    assert payload["document_type"] == "html"
    assert payload["text"] == "Policy\n\nUseful text.\n\nA | B"
    assert "ignore me" not in payload["text"]
    assert payload["metadata"]["headings"][0]["level"] == 1


def test_pdf_extractor_page_limit_and_text(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    import pypdf

    class Page:
        def __init__(self, text: str) -> None:
            self.text = text

        def extract_text(self) -> str:
            return self.text

    class Reader:
        def __init__(self, source: Any) -> None:
            self.pages = [Page("Page one"), Page("Page two")]

    monkeypatch.setattr(pypdf, "PdfReader", Reader)
    text, metadata = main._extract_pdf(b"fake")
    assert text == "Page one\n\nPage two"
    assert metadata == {"pages": 2}

    class LongReader:
        def __init__(self, source: Any) -> None:
            self.pages = [Page("x")] * 201

    monkeypatch.setattr(pypdf, "PdfReader", LongReader)
    with pytest.raises(ValueError, match="too many pages"):
        main._extract_pdf(b"fake")


def test_pdf_and_html_extraction_errors_are_sanitized(
    client: TestClient, monkeypatch: pytest.MonkeyPatch
) -> None:
    def broken(data: bytes) -> tuple[str, dict[str, Any]]:
        raise ValueError("internal parser detail")

    monkeypatch.setattr(main, "_extract_pdf", broken)
    monkeypatch.setattr(main, "_extract_html", broken)
    assert client.post("/extract", files={"file": ("x.pdf", b"x")}).status_code == 400
    assert client.post("/extract", files={"file": ("x.html", b"x")}).status_code == 400


@pytest.mark.parametrize(
    "payload",
    [
        b"not a zip",
        zipfile.ZipFile,
    ],
)
def test_malformed_docx_rejected(client: TestClient, payload: Any) -> None:
    if payload is zipfile.ZipFile:
        buffer = io.BytesIO()
        with zipfile.ZipFile(buffer, "w") as archive:
            archive.writestr("other.xml", "<x/>")
        data = buffer.getvalue()
    else:
        data = payload
    response = client.post("/extract", files={"file": ("bad.docx", data)})
    assert response.status_code == 400


def test_upload_size_and_extension_validation(client: TestClient) -> None:
    assert (
        client.post(
            "/extract",
            files={"file": ("large.txt", b"x" * (main.MAX_UPLOAD_BYTES + 1))},
        ).status_code
        == 413
    )
    assert (
        client.post("/extract", files={"file": ("sample.pdf", b"%PDF")}).status_code
        == 400
    )


def test_docx_archive_resource_limits(
    client: TestClient, monkeypatch: pytest.MonkeyPatch
) -> None:
    document = Document()
    document.add_paragraph("Text")
    buffer = io.BytesIO()
    document.save(buffer)
    monkeypatch.setattr(main, "MAX_DOCX_ENTRIES", 0)
    response = client.post(
        "/extract", files={"file": ("limited.docx", buffer.getvalue())}
    )
    assert response.status_code == 400


def test_docx_export_is_readable_and_sanitizes_filename(client: TestClient) -> None:
    response = client.post(
        "/export/docx",
        json={"text": "One\n\nCafé 😀", "title": 'Unsafe\\r\\n"name'},
    )
    assert response.status_code == 200
    assert response.headers["content-type"].startswith("application/vnd.openxmlformats")
    output = Document(io.BytesIO(response.content))
    assert [paragraph.text for paragraph in output.paragraphs] == [
        'Unsafe\\r\\n"name',
        "One",
        "Café 😀",
    ]
    assert "\r" not in response.headers["content-disposition"]


def test_docx_export_rejects_empty_and_oversize(client: TestClient) -> None:
    assert client.post("/export/docx", json={"text": ""}).status_code == 400
    assert client.post("/export/docx", json={"text": "x" * 100_001}).status_code == 413
