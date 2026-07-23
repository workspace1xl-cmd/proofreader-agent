from __future__ import annotations

import asyncio
import io
import json
import logging
import re
import zipfile
from contextlib import suppress
from pathlib import Path
from typing import Annotated, Any

from fastapi import FastAPI, File, HTTPException, Request, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel, Field

from app.cerebras_client import CEREBRAS_API_KEY, MODEL
from app.config import (
    MAX_CHARS,
    MAX_DOCX_ENTRIES,
    MAX_DOCX_UNCOMPRESSED_BYTES,
    MAX_UPLOAD_BYTES,
    PIPELINE_TIMEOUT_SECONDS,
    SSE_HEARTBEAT_SECONDS,
)
from app.pipeline import run_pipeline
from app.supabase_client import get_job, list_jobs, save_job, storage_mode

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s %(levelname)s %(name)s %(message)s",
)
logger = logging.getLogger(__name__)
ROOT = Path(__file__).resolve().parent.parent

app = FastAPI(title="Proofreading Agent", version="3.0.0")
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["GET", "POST", "HEAD", "OPTIONS"],
    allow_headers=["Content-Type"],
)


class ProofreadRequest(BaseModel):
    text: str
    document_type: str = "auto"
    metadata: dict[str, Any] = Field(default_factory=dict)


class ExportRequest(BaseModel):
    text: str
    title: str = "Proofread document"


def _validate_text(text: str) -> None:
    if not text or not text.strip():
        raise HTTPException(status_code=400, detail="Text must not be empty.")
    if len(text) > MAX_CHARS:
        raise HTTPException(
            status_code=413,
            detail=f"Text exceeds the {MAX_CHARS:,} character limit.",
        )


async def _run_with_timeout(
    text: str,
    progress: Any = None,
    agent_status: Any = None,
    document_type: str = "auto",
    metadata: dict[str, Any] | None = None,
) -> dict[str, Any]:
    try:
        async with asyncio.timeout(PIPELINE_TIMEOUT_SECONDS):
            return await run_pipeline(
                text,
                progress,
                agent_status,
                document_type,
                metadata,
            )
    except TimeoutError as exc:
        raise RuntimeError(
            "The review timed out. Please retry or split this document into sections."
        ) from exc


@app.get("/health")
def health() -> dict[str, Any]:
    return {
        "status": "ok",
        "version": app.version,
        "model": MODEL,
        "storage": storage_mode(),
        "cerebras_key_configured": bool(CEREBRAS_API_KEY),
        "max_chars": MAX_CHARS,
    }


@app.post("/proofread")
async def proofread_endpoint(req: ProofreadRequest) -> dict[str, Any]:
    _validate_text(req.text)
    try:
        result = await _run_with_timeout(
            req.text,
            document_type=req.document_type,
            metadata=req.metadata,
        )
    except RuntimeError as exc:
        raise HTTPException(status_code=503, detail=str(exc)) from exc
    saved = await asyncio.to_thread(save_job, req.text, result)
    return {"id": saved.get("id"), **result}


@app.post("/proofread/stream")
async def proofread_stream(
    req: ProofreadRequest, request: Request
) -> StreamingResponse:
    _validate_text(req.text)
    queue: asyncio.Queue[dict[str, Any] | None] = asyncio.Queue(maxsize=64)

    async def progress(done: int, total: int) -> None:
        await queue.put({"type": "progress", "done": done, "total": total})

    async def agent_status(key: str, label: str, state: str) -> None:
        await queue.put(
            {
                "type": "agent",
                "key": key,
                "label": label,
                "state": state,
            }
        )

    async def runner() -> None:
        try:
            result = await _run_with_timeout(
                req.text,
                progress,
                agent_status,
                req.document_type,
                req.metadata,
            )
            saved = await asyncio.to_thread(save_job, req.text, result)
            await queue.put(
                {"type": "result", "payload": {"id": saved.get("id"), **result}}
            )
        except asyncio.CancelledError:
            raise
        except Exception:
            logger.exception("Streaming review failed")
            await queue.put(
                {
                    "type": "error",
                    "detail": (
                        "The review could not be completed. Please retry in a moment."
                    ),
                }
            )
        finally:
            await queue.put(None)

    async def events() -> Any:
        task = asyncio.create_task(runner(), name="proofread-pipeline")
        event_id = 0
        try:
            yield "retry: 3000\n\n"
            while True:
                if await request.is_disconnected():
                    break
                try:
                    item = await asyncio.wait_for(
                        queue.get(), timeout=SSE_HEARTBEAT_SECONDS
                    )
                except TimeoutError:
                    yield ": keep-alive\n\n"
                    continue
                if item is None:
                    break
                event_id += 1
                payload = json.dumps(item, ensure_ascii=False)
                yield f"id: {event_id}\ndata: {payload}\n\n"
        except asyncio.CancelledError:
            raise
        finally:
            if not task.done():
                task.cancel()
            with suppress(asyncio.CancelledError):
                await task

    return StreamingResponse(
        events(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache, no-store",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        },
    )


def _validate_docx_archive(data: bytes) -> None:
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as archive:
            entries = archive.infolist()
            if len(entries) > MAX_DOCX_ENTRIES:
                raise ValueError("too many archive entries")
            total = sum(entry.file_size for entry in entries)
            if total > MAX_DOCX_UNCOMPRESSED_BYTES:
                raise ValueError("expanded document is too large")
            if not any(entry.filename == "word/document.xml" for entry in entries):
                raise ValueError("missing Word document content")
    except (zipfile.BadZipFile, OSError, ValueError) as exc:
        raise ValueError("invalid DOCX archive") from exc


def _extract_docx(data: bytes) -> tuple[str, dict[str, Any]]:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    _validate_docx_archive(data)
    document = Document(io.BytesIO(data))
    parts: list[str] = []
    headings: list[dict[str, Any]] = []
    tables: list[dict[str, Any]] = []
    offset = 0

    def append_part(value: str) -> int:
        nonlocal offset
        if parts:
            offset += 2
        start = offset
        parts.append(value)
        offset += len(value)
        return start

    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            paragraph_object = Paragraph(child, document)
            paragraph = paragraph_object.text
            if paragraph.strip():
                start = append_part(paragraph)
                style_name = str(getattr(paragraph_object.style, "name", "") or "")
                heading_match = re.match(r"Heading\s+(\d+)", style_name, re.I)
                if heading_match:
                    headings.append(
                        {
                            "text": paragraph,
                            "level": int(heading_match.group(1)),
                            "offset": start,
                        }
                    )
        elif child.tag == qn("w:tbl"):
            rows: list[str] = []
            column_counts: list[int] = []
            for row in Table(child, document).rows:
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                if any(cells):
                    rows.append("| " + " | ".join(cells) + " |")
                    column_counts.append(len(cells))
            if rows:
                table_text = "\n".join(rows)
                start = append_part(table_text)
                tables.append(
                    {
                        "offset": start,
                        "rows": len(rows),
                        "column_counts": column_counts,
                    }
                )
    return "\n\n".join(parts), {"headings": headings, "tables": tables}


def _extract_html(data: bytes) -> tuple[str, dict[str, Any]]:
    from bs4 import BeautifulSoup

    source = data.decode("utf-8-sig", errors="replace")
    soup = BeautifulSoup(source, "html.parser")
    for element in soup(["script", "style", "noscript"]):
        element.decompose()
    blocks: list[str] = []
    headings: list[dict[str, Any]] = []
    offset = 0
    for element in soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6", "p", "li", "tr"]):
        separator = " | " if element.name == "tr" else " "
        value = separator.join(element.stripped_strings)
        if not value:
            continue
        if blocks:
            offset += 2
        start = offset
        blocks.append(value)
        offset += len(value)
        if element.name and re.fullmatch(r"h[1-6]", element.name):
            headings.append(
                {"text": value, "level": int(element.name[1]), "offset": start}
            )
    text = "\n\n".join(blocks) or soup.get_text("\n", strip=True)
    return text, {"headings": headings}


def _extract_pdf(data: bytes) -> tuple[str, dict[str, Any]]:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    if len(reader.pages) > 200:
        raise ValueError("PDF has too many pages")
    pages = [str(page.extract_text() or "").strip() for page in reader.pages]
    text = "\n\n".join(page for page in pages if page)
    return text, {"pages": len(reader.pages)}


async def _read_limited(file: UploadFile) -> bytes:
    data = bytearray()
    while chunk := await file.read(64 * 1024):
        data.extend(chunk)
        if len(data) > MAX_UPLOAD_BYTES:
            raise HTTPException(status_code=413, detail="File too large (max 2 MB).")
    return bytes(data)


@app.post("/extract")
async def extract_endpoint(
    file: Annotated[UploadFile, File()],
) -> dict[str, Any]:
    name = (file.filename or "").lower()
    data = await _read_limited(file)
    if name.endswith(".docx"):
        try:
            text, metadata = await asyncio.to_thread(_extract_docx, data)
        except Exception as exc:
            logger.info("Rejected malformed DOCX %r: %s", file.filename, exc)
            raise HTTPException(
                status_code=400, detail="Could not read this .docx file."
            ) from exc
        document_type = "docx"
    elif name.endswith((".html", ".htm")):
        try:
            text, metadata = await asyncio.to_thread(_extract_html, data)
        except Exception as exc:
            raise HTTPException(
                status_code=400, detail="Could not read this HTML file."
            ) from exc
        document_type = "html"
    elif name.endswith(".pdf"):
        try:
            text, metadata = await asyncio.to_thread(_extract_pdf, data)
        except Exception as exc:
            raise HTTPException(
                status_code=400, detail="Could not read this PDF file."
            ) from exc
        document_type = "pdf"
    elif name.endswith((".txt", ".md")):
        text = data.decode("utf-8-sig", errors="replace")
        metadata = {}
        document_type = "markdown" if name.endswith(".md") else "txt"
    else:
        raise HTTPException(
            status_code=400,
            detail=("Unsupported file type. Upload TXT, Markdown, DOCX, HTML, or PDF."),
        )
    _validate_text(text)
    return {
        "filename": file.filename,
        "text": text,
        "chars": len(text),
        "document_type": document_type,
        "metadata": metadata,
    }


@app.post("/export/docx")
async def export_docx(req: ExportRequest) -> StreamingResponse:
    _validate_text(req.text)
    from docx import Document

    def build() -> io.BytesIO:
        document = Document()
        document.add_heading(req.title[:120], level=1)
        for paragraph in req.text.split("\n\n"):
            document.add_paragraph(paragraph)
        buffer = io.BytesIO()
        document.save(buffer)
        buffer.seek(0)
        return buffer

    buffer = await asyncio.to_thread(build)
    safe_name = (
        "".join(char for char in req.title if char.isalnum() or char in " -_")[:60]
        or "document"
    )
    return StreamingResponse(
        buffer,
        media_type=(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ),
        headers={"Content-Disposition": f'attachment; filename="{safe_name}.docx"'},
    )


@app.get("/jobs/{job_id}")
async def get_job_endpoint(job_id: str) -> dict[str, Any]:
    job = await asyncio.to_thread(get_job, job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found.")
    return job


@app.get("/jobs")
async def list_jobs_endpoint(limit: int = 20) -> list[dict[str, Any]]:
    return await asyncio.to_thread(list_jobs, min(max(limit, 1), 100))


app.mount("/static", StaticFiles(directory=ROOT / "static"), name="static")


@app.api_route("/", methods=["GET", "HEAD"])
def root() -> FileResponse:
    return FileResponse(ROOT / "static" / "index.html")
